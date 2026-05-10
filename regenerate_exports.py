#!/usr/bin/env python3
"""
Regenera Excel/Word usando rubric_config.yaml actual y puntajes manuales dados.
Uso:
  python regenerate_exports.py
  OUT_DIR=~/Downloads python regenerate_exports.py

Por defecto repite los puntajes del ejemplo que el usuario había exportado (66% con rúbrica vieja).
Con la rúbrica nueva (Anexo V) el porcentaje total cambia aunque los 0–4 sean iguales.
"""
from __future__ import annotations

import io
import os
import sys
from datetime import datetime
from pathlib import Path

import yaml
from docx import Document
from docx.shared import Pt
from openpyxl import Workbook

_APP_DIR = Path(__file__).resolve().parent


def weighted_score(scores: dict, weights: dict) -> float:
    total = sum(scores[s] * weights[s] for s in scores)
    max_total = sum(weights.values()) * 4
    return (total / max_total) * 100 if max_total > 0 else 0.0


def generate_excel(scores, percent, thresholds, label_fn):
    wb = Workbook()
    ws = wb.active
    ws.title = "Resultados"
    ws.append(["Criterio", "Puntaje (0–4)"])
    for k, v in scores.items():
        ws.append([label_fn(k), v])
    ws.append([])
    ws.append(["Puntaje total (%)", round(percent, 2)])
    if percent >= thresholds["aprobado"]:
        result = "Aprobado"
    elif percent >= thresholds["aprobado_obs"]:
        result = "Aprobado con observaciones"
    else:
        result = "No aprobado"
    ws.append(["Dictamen", result])
    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    return output.getvalue()


def generate_word(scores, percent, thresholds, nombre_proyecto: str, label_fn):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    base_title = "UCCuyo – Valoración de Informe de Avance"
    nombre_clean = (nombre_proyecto or "").strip()
    if nombre_clean:
        doc.add_heading(f'{base_title} "Del proyecto {nombre_clean}"', level=1)
    else:
        doc.add_heading(base_title, level=1)

    doc.add_paragraph(f"Fecha: {datetime.today().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph("")

    doc.add_heading("Resultados por criterio", level=2)
    table = doc.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text = "Criterio"
    hdr[1].text = "Puntaje (0–4)"
    for k, v in scores.items():
        row = table.add_row().cells
        row[0].text = label_fn(k)
        row[1].text = str(v)

    doc.add_paragraph(f"\nCumplimiento: {round(percent, 2)}%")

    if percent >= thresholds["aprobado"]:
        result = "Aprobado"
    elif percent >= thresholds["aprobado_obs"]:
        result = "Aprobado con observaciones"
    else:
        result = "No aprobado"

    doc.add_heading("Dictamen final", level=2)
    doc.add_paragraph(result)

    doc.add_heading("Observaciones del evaluador", level=2)
    doc.add_paragraph("." * 78)
    doc.add_paragraph("." * 78)
    doc.add_paragraph("." * 78)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def main() -> int:
    with open(_APP_DIR / "rubric_config.yaml", "r", encoding="utf-8") as f:
        config = yaml.safe_load(f)

    weights: dict = config["weights"]
    thresholds = config["thresholds"]
    labels: dict = config.get("labels") or {}

    def criterion_label(key: str) -> str:
        return labels.get(key, key.replace("_", " ").title())

    # Mismos puntajes 0–4 que el export anterior (impacto 0 → etica_normativa 0).
    manual_scores = {
        "identificacion": 4,
        "cronograma": 3,
        "objetivos": 2,
        "metodologia": 4,
        "resultados": 4,
        "formacion": 3,
        "gestion": 2,
        "dificultades": 1,
        "difusion": 1,
        "calidad_formal": 1,
        "etica_normativa": 0,
    }

    missing = set(weights) - set(manual_scores)
    extra = set(manual_scores) - set(weights)
    if missing or extra:
        print("Error: scores y weights no coinciden.", file=sys.stderr)
        print("  Faltan en manual_scores:", missing, file=sys.stderr)
        print("  Sobran en manual_scores:", extra, file=sys.stderr)
        return 1

    ordered_scores = {k: manual_scores[k] for k in weights}
    percent = weighted_score(ordered_scores, weights)

    out_dir = Path(os.environ.get("OUT_DIR", str(Path.home() / "Downloads"))).expanduser()
    out_dir.mkdir(parents=True, exist_ok=True)
    base = out_dir / "valoracion_informe_avance_rubrica_anexo_v"

    xlsx_bytes = generate_excel(ordered_scores, percent, thresholds, criterion_label)
    docx_bytes = generate_word(ordered_scores, percent, thresholds, "", criterion_label)

    xlsx_path = out_dir / f"{base.name}.xlsx"
    docx_path = out_dir / f"{base.name}.docx"
    xlsx_path.write_bytes(xlsx_bytes)
    docx_path.write_bytes(docx_bytes)

    print("Rúbrica:", _APP_DIR / "rubric_config.yaml")
    print("Puntaje total con ponderación Anexo V:", round(percent, 2), "%")
    print("Archivos:")
    print(" ", xlsx_path)
    print(" ", docx_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
