
# word_utils_fix.py
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import List, Optional

def add_full_text(doc: Document, text: Optional[str] = None):
    """Agrega texto completo preservando párrafos y saltos de línea.
    No acorta ni agrega '...'.
    """
    if not text:
        return
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    blocks = text.split('\n\n')
    for block in blocks:
        for line in block.split('\n'):
            doc.add_paragraph(line)
        doc.add_paragraph('')  # separación entre bloques

def add_table(doc: Document, headers, rows):
    if not headers or rows is None:
        return
    tbl = doc.add_table(rows=1, cols=len(headers))
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = str(h)
    for r in rows:
        cells = tbl.add_row().cells
        for i, h in enumerate(headers):
            cells[i].text = str(r.get(h, ""))
    return tbl

def export_informe_avance(path: str,
                          encabezado: str,
                          proyecto: str,
                          calificacion: str,
                          dictamen: str,
                          interpretacion: str = "",
                          observaciones: str = "",
                          tablas: dict = None):
    d = Document()
    p = d.add_paragraph(encabezado or "Universidad Católica de Cuyo — Secretaría de Investigación")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d.add_paragraph("Informe de Avance — Dictamen de Evaluación").alignment = WD_ALIGN_PARAGRAPH.CENTER
    d.add_paragraph("")

    if proyecto:
        d.add_paragraph(f"Proyecto: {proyecto}")
    if calificacion:
        d.add_paragraph(f"Resultado: {calificacion}")
    d.add_paragraph("")

    d.add_paragraph("Dictamen")
    add_full_text(d, dictamen)

    if interpretacion:
        d.add_paragraph("Interpretación")
        add_full_text(d, interpretacion)

    if observaciones:
        d.add_paragraph("Observaciones")
        add_full_text(d, observaciones)

    if tablas:
        for titulo, data in tablas.items():
            d.add_paragraph(titulo)
            headers = data.get("headers", [])
            rows = data.get("rows", [])
            add_table(d, headers, rows)

    d.save(path)
    return path
