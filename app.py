import streamlit as st
import pandas as pd
import numpy as np
import pdfplumber
import yaml
import io
from docx import Document
from docx.shared import Pt
from datetime import datetime
from openpyxl import Workbook
from pathlib import Path

_APP_DIR = Path(__file__).resolve().parent

# Mismo PNG en `assets/` (push a main). Fallback por URL cuando el archivo aún no está en el deploy.
_ESCUDO_REMOTE_URL = (
    "https://raw.githubusercontent.com/claudiomlarrea/valorador_informes_avances/"
    "main/assets/escudo_uccuyo.png"
)


def _resolve_escudo_path() -> Path | None:
    assets = _APP_DIR / "assets"
    if not assets.is_dir():
        return None
    for name in ("escudo_uccuyo.png", "escudo_uccuyo.jpg", "escudo_uccuyo.jpeg"):
        p = assets / name
        if p.is_file():
            return p
    return None


def _escudo_display_source() -> str:
    p = _resolve_escudo_path()
    return str(p) if p is not None else _ESCUDO_REMOTE_URL


st.set_page_config(layout="wide")
# ============================
# CONFIGURACIÓN
# ============================
with open("rubric_config.yaml", "r", encoding="utf-8") as f:
    config = yaml.safe_load(f)

weights = config["weights"]
thresholds = config["thresholds"]
keywords = config["keywords"]

# ============================
# FUNCIONES
# ============================
def extract_text(file):
    """Extrae texto desde PDF o DOCX"""
    if file.name.endswith(".pdf"):
        text = ""
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                text += (page.extract_text() or "") + "\n"
        return text
    elif file.name.endswith(".docx"):
        doc = Document(file)
        return "\n".join([p.text for p in doc.paragraphs])
    else:
        return ""

def auto_score(text, keywords_dict):
    """Calcula puntajes automáticos según palabras clave"""
    scores = {}
    text_low = (text or "").lower()
    for section, keys in keywords_dict.items():
        found = sum((k or "").lower() in text_low for k in keys)
        scores[section] = min(4, found)
    return scores

def weighted_score(scores, weights):
    """Calcula el puntaje total ponderado (%) a partir de puntajes 0–4"""
    total = sum(scores[s] * weights[s] for s in scores)
    max_total = sum(weights.values()) * 4
    percent = (total / max_total) * 100 if max_total > 0 else 0.0
    return percent

def generate_excel(scores, percent, thresholds):
    """Genera archivo Excel con resultados"""
    wb = Workbook()
    ws = wb.active
    ws.title = "Resultados"
    ws.append(["Criterio", "Puntaje (0–4)"])
    for k, v in scores.items():
        ws.append([k, v])
    ws.append([])
    ws.append(["Puntaje total (%)", round(percent, 2)])
    if percent >= thresholds["aprobado"]:
        result = "Aprobado"
    elif percent >= thresholds["aprobado_obs"]:
        result = "Aprobado con observaciones"
    else:
        result = "No aprobado"
    ws.append(["Dictamen", result])
    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    return output

def generate_word(scores, percent, thresholds, nombre_proyecto=""):
    """Genera dictamen Word incluyendo el nombre del proyecto"""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    # Encabezado
    base_title = "UCCuyo – Valoración de Informe de Avance"
    nombre_clean = (nombre_proyecto or "").strip()
    if nombre_clean:
        doc.add_heading(f'{base_title} "Del proyecto {nombre_clean}"', level=1)
    else:
        doc.add_heading(base_title, level=1)

    doc.add_paragraph(f"Fecha: {datetime.today().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph("")

    # Puntajes por criterio
    doc.add_heading("Resultados por criterio", level=2)
    table = doc.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text = "Criterio"
    hdr[1].text = "Puntaje (0–4)"
    for k, v in scores.items():
        row = table.add_row().cells
        row[0].text = k.replace("_", " ").capitalize()
        row[1].text = str(v)

    percent_text = f"\nCumplimiento: {round(percent, 2)}%"
    doc.add_paragraph(percent_text)

    # Dictamen final
    if percent >= thresholds["aprobado"]:
        result = "Aprobado"
    elif percent >= thresholds["aprobado_obs"]:
        result = "Aprobado con observaciones"
    else:
        result = "No aprobado"

    doc.add_heading("Dictamen final", level=2)
    doc.add_paragraph(result)

    # Observaciones
    doc.add_heading("Observaciones del evaluador", level=2)
    doc.add_paragraph("..............................................................................")
    doc.add_paragraph("..............................................................................")
    doc.add_paragraph("..............................................................................")

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# ============================
# INTERFAZ STREAMLIT
# ============================
st.markdown(
    """
<style>
:root {
    --ucc-green: #00664d;
    --ucc-green-dark: #00523e;
    --ucc-accent: #28a745;
    --ucc-page-bg: #f8f9fa;
    --ucc-sidebar-bg: #262730;
    --ucc-text: #262730;
    --ucc-heading-card: #2c3838;
    --ucc-lead-muted: #5f6b6f;
}

.stApp {
    background-color: var(--ucc-page-bg);
}

/* Chrome superior de Streamlit: evitar banda oscura y solape con marca UCCuyo */
header[data-testid="stHeader"] {
    background: var(--ucc-page-bg) !important;
    border-bottom: 1px solid rgba(0, 0, 0, 0.06);
}
div[data-testid="stDecoration"] {
    height: 3px !important;
    margin-top: env(safe-area-inset-top, 0);
    background: linear-gradient(
        90deg,
        var(--ucc-green-dark) 0%,
        var(--ucc-green) 50%,
        var(--ucc-green-dark) 100%
    ) !important;
}

.block-container {
    padding-top: 2rem !important;
    padding-left: calc(1rem + env(safe-area-inset-left, 0px)) !important;
    padding-right: calc(1rem + env(safe-area-inset-right, 0px)) !important;
}

section[data-testid="stSidebar"] {
    background-color: var(--ucc-sidebar-bg);
}
[data-testid="stSidebar"] [data-testid="stMarkdown"],
[data-testid="stSidebar"] span,
[data-testid="stSidebar"] label {
    color: rgba(255, 255, 255, 0.92);
}

.ucc-banner {
    background: var(--ucc-green);
    border-radius: 12px;
    padding: 1.35rem 1.65rem;
    display: flex;
    flex-direction: column;
    justify-content: center;
    box-sizing: border-box;
    min-height: 120px;
}
.header-uccuyo h1.ucc-banner-heading,
.header-uccuyo h2.ucc-banner-heading,
.header-uccuyo h3.ucc-banner-heading {
    color: #ffffff !important;
    margin: 0;
    line-height: 1.2;
    font-family: "Source Sans Pro", ui-sans-serif, system-ui, sans-serif;
}
.header-uccuyo h1.ucc-banner-heading {
    font-size: clamp(1.35rem, 2.8vw, 1.95rem);
    font-weight: 700;
}
.header-uccuyo h2.ucc-banner-heading {
    margin-top: 0.55rem !important;
    font-size: clamp(1rem, 2vw, 1.25rem);
    font-weight: 500;
}
.header-uccuyo h3.ucc-banner-heading {
    margin-top: 0.35rem !important;
    font-size: clamp(0.85rem, 1.4vw, 1rem);
    font-weight: 400;
    color: rgba(255, 255, 255, 0.92) !important;
}

.block-container > div:first-child div[data-testid="stHorizontalBlock"] {
    margin-bottom: 1.35rem;
    align-items: stretch;
}
.block-container > div:first-child div[data-testid="stHorizontalBlock"] > div[data-testid="column"]:first-child {
    display: flex;
    flex-direction: column;
    justify-content: center;
}
.block-container > div:first-child div[data-testid="stHorizontalBlock"] img {
    border: 4px solid var(--ucc-green);
    border-radius: 10px;
    background: #fff;
    display: block;
}

h1:not(.ucc-banner-heading):not(.uc-card-main-title),
h2:not(.ucc-banner-heading),
h3:not(.ucc-banner-heading),
h4 {
    color: var(--ucc-green-dark) !important;
}

/* Tarjeta intro (misma línea visual que otros sistemas institucionales Streamlit) */
.ucc-intro-card {
    background: #ffffff;
    border-radius: 14px;
    padding: 1.75rem 2rem;
    margin-bottom: 1.65rem;
    box-shadow:
        0 8px 28px rgba(0, 0, 0, 0.07),
        0 1px 3px rgba(0, 0, 0, 0.04);
}
.ucc-intro-card h1.uc-card-main-title {
    color: var(--ucc-heading-card) !important;
    margin: 0 0 0.75rem 0 !important;
    font-size: clamp(1.3rem, 2.8vw, 1.85rem);
    font-weight: 700;
    line-height: 1.25;
    font-family: "Source Sans Pro", ui-sans-serif, system-ui, sans-serif;
}
.ucc-intro-card p.uc-card-lead {
    color: var(--ucc-lead-muted) !important;
    margin: 0 !important;
    line-height: 1.6;
    font-size: 1.02rem;
}

p:not(.ucc-banner-heading):not(.uc-card-lead),
label {
    color: var(--ucc-text) !important;
}

/* Controles densos tipo app de prácticos */
[data-testid="stTextInput"] input,
[data-testid="stNumberInput"] input {
    border-radius: 12px !important;
    border-width: 0 !important;
}
[data-baseweb="select"] > div:first-child {
    border-radius: 12px !important;
}

[data-testid="stFileUploader"] {
    background: linear-gradient(#fcfdfd, #f4f6f7) !important;
    border-radius: 14px !important;
    padding: 1.15rem 1.25rem !important;
    border: 2px dashed rgba(0, 82, 62, 0.28) !important;
    box-sizing: border-box;
}
[data-testid="stFileUploader"] button {
    background-color: var(--ucc-green) !important;
    color: white !important;
    border-radius: 8px;
    border: none;
}
.stButton > button {
    background-color: var(--ucc-green) !important;
    color: white !important;
    border-radius: 8px;
    border: none;
    font-weight: 600;
}
.stButton > button:hover {
    background-color: var(--ucc-green-dark) !important;
    border-color: transparent !important;
}
[data-testid="stDownloadButton"] button {
    background-color: var(--ucc-green) !important;
    color: white !important;
    border-radius: 8px;
    border: none;
    font-weight: 600;
}
[data-testid="stDownloadButton"] button:hover {
    background-color: var(--ucc-green-dark) !important;
    border-color: transparent !important;
}
div[data-testid="stAlert"] {
    border-radius: 10px;
}
[data-baseweb="slider"] {
    color: var(--ucc-green);
}
.stButton button span,
[data-testid="stDownloadButton"] button span {
    color: white !important;
}
.stButton > button,
.stButton > button * {
    color: white !important;
}
[data-testid="stDownloadButton"] button,
[data-testid="stDownloadButton"] button * {
    color: white !important;
}
[data-testid="stFileUploader"] button span,
[data-testid="stFileUploader"] button div,
[data-testid="stFileUploader"] button p {
    color: white !important;
}

.stSlider label,
[data-testid="stTextInput"] label,
[data-testid="stFileUploader"] label {
    position: relative;
    padding-left: 1rem;
}
.stSlider label::before,
[data-testid="stTextInput"] label::before,
[data-testid="stFileUploader"] label::before {
    content: "";
    position: absolute;
    left: 0;
    top: 0.45rem;
    width: 9px;
    height: 9px;
    border-radius: 50%;
    background: var(--ucc-accent);
}
</style>
""",
    unsafe_allow_html=True,
)

_banner_html = """<div class="ucc-banner header-uccuyo">
<h1 class="ucc-banner-heading">Universidad Católica de Cuyo</h1>
<h2 class="ucc-banner-heading">Secretaría de Investigación</h2>
<h3 class="ucc-banner-heading">Consejo de Investigación</h3>
</div>"""

_brand_logo_col, _brand_banner_col = st.columns([1, 6], gap="medium")
with _brand_logo_col:
    st.image(_escudo_display_source(), width=118, use_container_width=False)
with _brand_banner_col:
    st.markdown(_banner_html, unsafe_allow_html=True)


st.markdown(
    """
<div class="ucc-intro-card">
<h1 class="uc-card-main-title">Valorador de Informes de Avance</h1>
<p class="uc-card-lead">Subí un informe de avance (PDF o DOCX) para evaluarlo automáticamente según la rúbrica institucional.</p>
</div>
""",
    unsafe_allow_html=True,
)

uploaded_file = st.file_uploader("Cargar archivo", type=["pdf", "docx"])

if uploaded_file:
    text = extract_text(uploaded_file)

    with st.expander("Ver texto extraído"):
        st.text_area("Texto completo", text, height=300)

    # --- Evaluación automática (referencia) ---
    st.subheader("Evaluación automática")
    auto_scores = auto_score(text, keywords)
    df = pd.DataFrame(auto_scores.items(), columns=["Criterio", "Puntaje (0–4)"])
    st.dataframe(df, use_container_width=True)

    auto_percent = weighted_score(auto_scores, weights)
    st.metric(label="Puntaje automático inicial (%)", value=round(auto_percent, 2))

    # --- Ajuste manual ---
    st.subheader("Ajuste manual (opcional)")
    manual_scores = {}
    for k in auto_scores.keys():
        manual_scores[k] = st.slider(
            f"{k.replace('_',' ').capitalize()}",
            0,
            4,
            int(auto_scores[k]),
        )

    # Puntaje total AJUSTADO (este es el que importa)
    adjusted_percent = weighted_score(manual_scores, weights)
    st.metric(label="Puntaje total ajustado (%)", value=round(adjusted_percent, 2))

    # Dictamen con ajuste manual
    if adjusted_percent >= thresholds["aprobado"]:
        result = "✅ Aprobado"
    elif adjusted_percent >= thresholds["aprobado_obs"]:
        result = "⚠️ Aprobado con observaciones"
    else:
        result = "❌ No aprobado"
    st.success(f"Dictamen (con ajuste manual): {result}")

    # Nombre del proyecto para el Word
    nombre_proyecto = st.text_input("Nombre del proyecto (aparecerá en el Word):", "")

    # Generar informes SIEMPRE con los valores ajustados
    if st.button("Generar informes"):
        final_percent = adjusted_percent
        excel_file = generate_excel(manual_scores, final_percent, thresholds)
        word_file = generate_word(manual_scores, final_percent, thresholds, nombre_proyecto)

        st.download_button(
            "⬇️ Descargar Excel",
            excel_file,
            file_name="valoracion_informe_avance.xlsx",
        )
        st.download_button(
            "⬇️ Descargar Word",
            word_file,
            file_name="valoracion_informe_avance.docx",
        )

        st.success("Informe generado con los puntajes ajustados manualmente.")
