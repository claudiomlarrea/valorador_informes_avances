# export_fix.py
# Genera el DOCX en bytes e incluye opcionalmente el nombre del proyecto en el título.
from io import BytesIO
from datetime import datetime
from docx import Document
from docx.shared import Pt


def export_word_dictamen(
    resultados: dict,
    cumplimiento,
    dictamen_texto: str,
    categoria: str = "",
    nombre_proyecto: str | None = None,
    fecha: str | None = None,
) -> bytes:
    """
    Crea el informe Word con encabezado institucional y devuelve BYTES (para st.download_button).

    - resultados: dict {"Identificacion":4, "Cronograma":3, ...}
    - cumplimiento: porcentaje (float o str)
    - dictamen_texto: 'Aprobado', 'Aprobado con observaciones' o 'No aprobado'
    - categoria: opcional (puede ir vacío)
    - nombre_proyecto: si viene, se muestra como "Del proyecto …" en el título
    - fecha: opcional (si no, se usa la actual)
    """
    if fecha is None:
        fecha = datetime.now().strftime("%Y-%m-%d %H:%M")

    base_titulo = "UCCuyo – Valoración de Informe de Avance"
    if nombre_proyecto and str(nombre_proyecto).strip():
        titulo = f'{base_titulo} "Del proyecto {str(nombre_proyecto).strip()}"'
    else:
        titulo = base_titulo

    doc = Document()

    # Título (con fallback por si el estilo no existe)
    try:
        p = doc.add_paragraph(titulo)
        p.style = "Title"
        p.runs[0].font.size = Pt(14)
    except Exception:
        p = doc.add_paragraph(titulo)
        p.runs[0].font.size = Pt(14)

    doc.add_paragraph(f"Fecha: {fecha}")
    doc.add_paragraph("")  # espacio

    # Resultados por criterio
    doc.add_paragraph("Resultados por criterio")
    table = doc.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text = "Criterio"
    hdr[1].text = "Puntaje (0–4)"

    for criterio, puntaje in resultados.items():
        row = table.add_row().cells
        row[0].text = str(criterio)
        row[1].text = str(puntaje)

    # Cumplimiento y dictamen
    try:
        cumpl_txt = f"{float(cumplimiento):.1f}%"
    except Exception:
        cumpl_txt = str(cumplimiento)

    doc.add_paragraph(f"\nCumplimiento: {cumpl_txt}")
    doc.add_paragraph("\nDictamen final")
    if categoria:
        doc.add_paragraph(str(categoria))
    if dictamen_texto:
        doc.add_paragraph(str(dictamen_texto))

    # Observaciones
    doc.add_paragraph(
        "\nObservaciones del evaluador\n"
        + "..............................................................................\n"
        + "..............................................................................\n"
        + ".............................................................................."
    )

    # A bytes
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# -------- Compatibilidad hacia atrás (si tu app llamaba export_word(...)) --------
def export_word(resultados, cumplimiento, dictamen_texto, categoria=""):
    """
    Wrapper para no romper apps antiguas. NO agrega nombre del proyecto.
    Si querés nombre de proyecto, usá export_word_dictamen(..., nombre_proyecto="...")
    """
    return export_word_dictamen(
        resultados=resultados,
        cumplimiento=cumplimiento,
        dictamen_texto=dictamen_texto,
        categoria=categoria,
        nombre_proyecto=None,  # mantenemos comportamiento previo
    )
